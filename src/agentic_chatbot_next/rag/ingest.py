from __future__ import annotations

import datetime as dt
import hashlib
import logging
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, List, Sequence

from langchain_core.documents import Document

from agentic_chatbot_next.config import Settings
from agentic_chatbot_next.persistence.postgres.collections import (
    COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES,
    COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS,
    normalize_collection_maintenance_policy,
)
from agentic_chatbot_next.persistence.postgres.requirements import RequirementStatementRecord
from agentic_chatbot_next.rag.clause_splitter import clause_split
from agentic_chatbot_next.rag.ocr import IMAGE_SUFFIXES, load_image_documents, load_pdf_documents_with_ocr
from agentic_chatbot_next.rag.requirements import (
    build_requirement_statement_records,
    supports_requirements_extraction,
)
from agentic_chatbot_next.rag.metadata_extractor import (
    INDEX_METADATA_VERSION,
    DocumentIndexMetadata,
    build_chunk_index_metadata,
    build_document_index_metadata,
    normalize_metadata_profile,
    summarize_index_metadata,
)
from agentic_chatbot_next.rag.corpus_metadata import enrich_corpus_metadata
from agentic_chatbot_next.rag.parser_adapters import ParsedDocumentBundle, load_documents_with_parsers
from agentic_chatbot_next.rag.structure_detector import (
    PROCESS_FLOW_PATTERN,
    REQUIREMENT_PATTERN,
    StructureAnalysis,
    detect_structure,
)
from agentic_chatbot_next.rag.status_workbooks import profile_workbook
from agentic_chatbot_next.rag.workbook_loader import load_workbook_documents
from agentic_chatbot_next.persistence.postgres.chunks import ChunkRecord
from agentic_chatbot_next.persistence.postgres.documents import DocumentRecord
from agentic_chatbot_next.rag.stores import KnowledgeStores, make_doc_id

logger = logging.getLogger(__name__)
_REPO_ROOT = Path(__file__).resolve().parents[3]
_REPO_PATH_ALIASES = (_REPO_ROOT, Path("/app"))


@dataclass(frozen=True)
class KBCoverageStatus:
    tenant_id: str
    collection_id: str
    configured_source_paths: tuple[str, ...]
    missing_source_paths: tuple[str, ...]
    indexed_source_paths: tuple[str, ...]
    indexed_doc_count: int
    sync_attempted: bool = False
    sync_error: str = ""
    synced_doc_ids: tuple[str, ...] = ()

    @property
    def ready(self) -> bool:
        return not self.sync_error and not self.missing_source_paths

    @property
    def status(self) -> str:
        return "ready" if self.ready else "not_ready"

    @property
    def reason(self) -> str:
        if self.sync_error:
            return "kb_sync_failed" if self.sync_attempted else "kb_status_check_failed"
        if self.missing_source_paths:
            return "kb_coverage_missing"
        return "ready"

    @property
    def suggested_fix(self) -> str:
        return f"python run.py sync-kb --collection-id {self.collection_id}"

    def to_dict(self) -> dict[str, Any]:
        return {
            "status": self.status,
            "reason": self.reason,
            "tenant_id": self.tenant_id,
            "collection_id": self.collection_id,
            "configured_source_count": len(self.configured_source_paths),
            "indexed_doc_count": self.indexed_doc_count,
            "missing_sources": list(self.missing_source_paths),
            "sync_attempted": self.sync_attempted,
            "sync_error": self.sync_error,
            "suggested_fix": self.suggested_fix,
        }


@dataclass(frozen=True)
class CollectionReadinessStatus:
    tenant_id: str
    collection_id: str
    maintenance_policy: str
    document_count: int = 0
    ready: bool = False
    reason: str = "empty_collection"
    configured_source_paths: tuple[str, ...] = ()
    missing_source_paths: tuple[str, ...] = ()
    indexed_source_paths: tuple[str, ...] = ()
    sync_error: str = ""
    suggested_fix: str = ""

    @property
    def indexed_doc_count(self) -> int:
        return self.document_count

    def to_dict(self) -> dict[str, Any]:
        return {
            "tenant_id": self.tenant_id,
            "collection_id": self.collection_id,
            "maintenance_policy": self.maintenance_policy,
            "document_count": self.document_count,
            "indexed_doc_count": self.document_count,
            "ready": self.ready,
            "reason": self.reason,
            "configured_source_paths": list(self.configured_source_paths),
            "missing_source_paths": list(self.missing_source_paths),
            "indexed_source_paths": list(self.indexed_source_paths),
            "sync_error": self.sync_error,
            "suggested_fix": self.suggested_fix,
        }


@dataclass(frozen=True)
class KBDocumentVersion:
    doc_id: str
    title: str
    source_type: str
    source_path: str
    collection_id: str
    content_hash: str
    ingested_at: str
    num_chunks: int = 0
    file_type: str = ""
    doc_structure_type: str = "general"
    active: bool = False
    version_ordinal: int = 1
    superseded_at: str = ""
    parser_chain: tuple[str, ...] = ()
    extraction_status: str = "success"
    extraction_error: str = ""
    metadata_confidence: float = 0.5
    lifecycle_phase: str = ""
    doc_type: str = ""
    program_entities: tuple[str, ...] = ()
    signal_summary: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "source_type": self.source_type,
            "source_path": self.source_path,
            "collection_id": self.collection_id,
            "content_hash": self.content_hash,
            "ingested_at": self.ingested_at,
            "num_chunks": self.num_chunks,
            "file_type": self.file_type,
            "doc_structure_type": self.doc_structure_type,
            "active": self.active,
            "version_ordinal": self.version_ordinal,
            "superseded_at": self.superseded_at,
            "parser_chain": list(self.parser_chain),
            "extraction_status": self.extraction_status,
            "extraction_error": self.extraction_error,
            "metadata_confidence": self.metadata_confidence,
            "lifecycle_phase": self.lifecycle_phase,
            "doc_type": self.doc_type,
            "program_entities": list(self.program_entities),
            "signal_summary": dict(self.signal_summary),
        }


@dataclass(frozen=True)
class KBSourceHealthGroup:
    source_identity: str
    title: str
    source_type: str
    collection_id: str
    configured_source_path: str = ""
    active_doc_id: str = ""
    active_content_hash: str = ""
    active_ingested_at: str = ""
    active_source_path: str = ""
    current_file_hash: str = ""
    source_exists: bool = False
    content_drift: bool = False
    duplicate_doc_ids: tuple[str, ...] = ()
    stale_version_doc_ids: tuple[str, ...] = ()
    extraction_failure_doc_ids: tuple[str, ...] = ()
    low_confidence_doc_ids: tuple[str, ...] = ()
    missing_source_doc_ids: tuple[str, ...] = ()
    parser_warnings: tuple[str, ...] = ()
    records: tuple[KBDocumentVersion, ...] = ()

    @property
    def status(self) -> str:
        if self.extraction_failure_doc_ids:
            return "extraction_failed"
        if self.low_confidence_doc_ids:
            return "low_confidence_metadata"
        if self.missing_source_doc_ids:
            return "missing_source"
        if self.duplicate_doc_ids:
            return "duplicate"
        if self.content_drift:
            return "content_drift"
        if self.parser_warnings:
            return "parser_warnings"
        return "healthy"

    def to_dict(self) -> dict[str, Any]:
        return {
            "source_identity": self.source_identity,
            "title": self.title,
            "source_type": self.source_type,
            "collection_id": self.collection_id,
            "configured_source_path": self.configured_source_path,
            "active_doc_id": self.active_doc_id,
            "active_content_hash": self.active_content_hash,
            "active_ingested_at": self.active_ingested_at,
            "active_source_path": self.active_source_path,
            "current_file_hash": self.current_file_hash,
            "source_exists": self.source_exists,
            "content_drift": self.content_drift,
            "duplicate_doc_ids": list(self.duplicate_doc_ids),
            "stale_version_doc_ids": list(self.stale_version_doc_ids),
            "extraction_failure_doc_ids": list(self.extraction_failure_doc_ids),
            "low_confidence_doc_ids": list(self.low_confidence_doc_ids),
            "missing_source_doc_ids": list(self.missing_source_doc_ids),
            "parser_warnings": list(self.parser_warnings),
            "status": self.status,
            "records": [item.to_dict() for item in self.records],
        }


@dataclass(frozen=True)
class KBCorpusHealthReport:
    tenant_id: str
    collection_id: str
    configured_source_paths: tuple[str, ...]
    missing_source_paths: tuple[str, ...]
    indexed_doc_count: int
    active_doc_count: int
    maintenance_policy: str = COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES
    duplicate_groups: tuple[KBSourceHealthGroup, ...] = ()
    drifted_groups: tuple[KBSourceHealthGroup, ...] = ()
    source_groups: tuple[KBSourceHealthGroup, ...] = ()
    sync_error: str = ""
    stale_version_count: int = 0
    extraction_failure_count: int = 0
    low_confidence_metadata_count: int = 0
    missing_source_doc_count: int = 0
    parser_warning_count: int = 0
    parser_counts: dict[str, int] = field(default_factory=dict)
    doc_type_counts: dict[str, int] = field(default_factory=dict)
    lifecycle_counts: dict[str, int] = field(default_factory=dict)
    metadata_confidence_distribution: dict[str, int] = field(default_factory=dict)

    @property
    def ready(self) -> bool:
        return (
            not self.sync_error
            and not self.missing_source_paths
            and not self.duplicate_groups
            and not self.drifted_groups
            and self.extraction_failure_count == 0
            and self.low_confidence_metadata_count == 0
            and self.missing_source_doc_count == 0
        )

    @property
    def reason(self) -> str:
        if self.sync_error:
            if self.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
                return "kb_status_check_failed"
            return "collection_status_check_failed"
        if self.duplicate_groups:
            if self.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
                return "kb_duplicate_docs"
            return "collection_duplicate_docs"
        if self.extraction_failure_count:
            return "collection_extraction_failures"
        if self.low_confidence_metadata_count:
            return "collection_low_confidence_metadata"
        if self.missing_source_doc_count:
            return "collection_source_missing"
        if self.drifted_groups:
            if self.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
                return "kb_content_drift"
            return "collection_content_drift"
        if self.missing_source_paths:
            if self.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
                return "kb_coverage_missing"
            return "collection_source_missing"
        return "ready"

    @property
    def suggested_fix(self) -> str:
        if self.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
            return f"python run.py repair-kb --collection-id {self.collection_id}"
        return (
            f"Use the collection repair action to prune duplicate or drifted documents in "
            f"'{self.collection_id}'."
        )

    def to_dict(self) -> dict[str, Any]:
        return {
            "status": "ready" if self.ready else "not_ready",
            "reason": self.reason,
            "tenant_id": self.tenant_id,
            "collection_id": self.collection_id,
            "maintenance_policy": self.maintenance_policy,
            "configured_source_count": len(self.configured_source_paths),
            "indexed_doc_count": self.indexed_doc_count,
            "active_doc_count": self.active_doc_count,
            "missing_sources": list(self.missing_source_paths),
            "duplicate_group_count": len(self.duplicate_groups),
            "content_drift_count": len(self.drifted_groups),
            "stale_version_count": self.stale_version_count,
            "extraction_failure_count": self.extraction_failure_count,
            "low_confidence_metadata_count": self.low_confidence_metadata_count,
            "missing_source_doc_count": self.missing_source_doc_count,
            "parser_warning_count": self.parser_warning_count,
            "parser_counts": dict(self.parser_counts),
            "doc_type_counts": dict(self.doc_type_counts),
            "lifecycle_counts": dict(self.lifecycle_counts),
            "metadata_confidence_distribution": dict(self.metadata_confidence_distribution),
            "duplicate_groups": [item.to_dict() for item in self.duplicate_groups],
            "drifted_groups": [item.to_dict() for item in self.drifted_groups],
            "source_groups": [item.to_dict() for item in self.source_groups],
            "sync_error": self.sync_error,
            "suggested_fix": self.suggested_fix,
        }


@dataclass(frozen=True)
class KBRepairResult:
    tenant_id: str
    collection_id: str
    deleted_doc_ids: tuple[str, ...] = ()
    reindexed_doc_ids: tuple[str, ...] = ()
    ingested_missing_doc_ids: tuple[str, ...] = ()
    unresolved_paths: tuple[str, ...] = ()
    health_before: KBCorpusHealthReport | None = None
    health_after: KBCorpusHealthReport | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "tenant_id": self.tenant_id,
            "collection_id": self.collection_id,
            "deleted_doc_ids": list(self.deleted_doc_ids),
            "reindexed_doc_ids": list(self.reindexed_doc_ids),
            "ingested_missing_doc_ids": list(self.ingested_missing_doc_ids),
            "unresolved_paths": list(self.unresolved_paths),
            "health_before": self.health_before.to_dict() if self.health_before else None,
            "health_after": self.health_after.to_dict() if self.health_after else None,
        }


@dataclass(frozen=True)
class RequirementBackfillResult:
    tenant_id: str
    collection_id: str
    source_type: str
    processed_doc_ids: tuple[str, ...] = ()
    statement_count: int = 0
    unsupported_documents: tuple[dict[str, str], ...] = ()

    def to_dict(self) -> dict[str, Any]:
        return {
            "tenant_id": self.tenant_id,
            "collection_id": self.collection_id,
            "source_type": self.source_type,
            "processed_doc_ids": list(self.processed_doc_ids),
            "processed_doc_count": len(self.processed_doc_ids),
            "statement_count": self.statement_count,
            "unsupported_documents": [dict(item) for item in self.unsupported_documents],
        }


def iter_kb_source_paths(settings: Settings) -> List[Path]:
    roots: List[Path] = []
    kb_dir = getattr(settings, "kb_dir", None)
    if kb_dir:
        roots.append(Path(kb_dir))
    roots.extend(Path(path) for path in getattr(settings, "kb_extra_dirs", ()) if path)
    paths: List[Path] = []
    seen: set[Path] = set()
    for root in roots:
        if not root.exists() or not root.is_dir():
            continue
        for path in sorted(root.glob("*")):
            if not path.is_file():
                continue
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            paths.append(resolved)
    return paths


def _normalize_path_text(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text.startswith("repo://"):
        return text
    try:
        resolved = Path(text).expanduser().resolve(strict=False)
    except Exception:
        return text.replace("\\", "/")
    for root in _REPO_PATH_ALIASES:
        try:
            relative = resolved.relative_to(root.resolve(strict=False))
        except Exception:
            continue
        return f"repo://{relative.as_posix()}"
    return str(resolved)


def canonicalize_local_source_path(value: str | Path) -> str:
    return _normalize_path_text(str(value or ""))


def _looks_like_path_identity(value: str) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    normalized = text.replace("\\", "/")
    if normalized.startswith("repo://"):
        return True
    if normalized.startswith(("/", "./", "../", "~/")):
        return True
    if len(text) >= 3 and text[1] == ":" and text[2] in {"/", "\\"} and text[0].isalpha():
        return True
    if "://" in normalized and not normalized.startswith("repo://"):
        return False
    return "/" in normalized


def _normalize_source_identity_text(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text.startswith("path:"):
        return f"path:{_normalize_path_text(text[len('path:') :])}"
    if _looks_like_path_identity(text):
        return f"path:{_normalize_path_text(text)}"
    return text


def _normalize_title(value: str) -> str:
    return Path(str(value or "").strip()).name.casefold()


def _list_collection_documents(
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str,
) -> list[Any]:
    list_documents = getattr(getattr(stores, "doc_store", None), "list_documents", None)
    if not callable(list_documents):
        return []
    return list(
        list_documents(
            source_type="",
            tenant_id=tenant_id,
            collection_id=collection_id,
        )
    )


def _collection_document_summary(
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str,
    records: Sequence[Any] | None = None,
) -> dict[str, Any]:
    get_collection_summary = getattr(getattr(stores, "doc_store", None), "get_collection_summary", None)
    if callable(get_collection_summary):
        summary = get_collection_summary(collection_id, tenant_id=tenant_id)
        if summary:
            return dict(summary)
    documents = list(records or _list_collection_documents(stores, tenant_id=tenant_id, collection_id=collection_id))
    latest_ingested_at = ""
    source_type_counts: dict[str, int] = {}
    for record in documents:
        source_type = str(getattr(record, "source_type", "") or "unknown").strip() or "unknown"
        source_type_counts[source_type] = int(source_type_counts.get(source_type, 0)) + 1
        ingested_at = str(getattr(record, "ingested_at", "") or "")
        if ingested_at and ingested_at > latest_ingested_at:
            latest_ingested_at = ingested_at
    return {
        "collection_id": collection_id,
        "document_count": len(documents),
        "latest_ingested_at": latest_ingested_at,
        "source_type_counts": source_type_counts,
    }


def infer_collection_maintenance_policy(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str,
    collection_record: Any | None = None,
    collection_summary: dict[str, Any] | None = None,
    documents: Sequence[Any] | None = None,
) -> str:
    explicit_policy = normalize_collection_maintenance_policy(
        str(getattr(collection_record, "maintenance_policy", "") or ""),
        default="",
    )
    if explicit_policy:
        return explicit_policy

    default_collection_id = str(getattr(settings, "default_collection_id", "default") or "default").strip() or "default"
    if collection_id == default_collection_id:
        return COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES

    summary = dict(collection_summary or {})
    source_type_counts = {
        str(key or "").strip().lower(): int(value or 0)
        for key, value in dict(summary.get("source_type_counts") or {}).items()
        if str(key or "").strip()
    }
    if not source_type_counts and documents is not None:
        for record in documents:
            source_type = str(getattr(record, "source_type", "") or "unknown").strip().lower() or "unknown"
            source_type_counts[source_type] = int(source_type_counts.get(source_type, 0)) + 1
    if not source_type_counts and getattr(stores, "doc_store", None) is not None:
        fallback_summary = _collection_document_summary(
            stores,
            tenant_id=tenant_id,
            collection_id=collection_id,
            records=documents,
        )
        source_type_counts = {
            str(key or "").strip().lower(): int(value or 0)
            for key, value in dict(fallback_summary.get("source_type_counts") or {}).items()
            if str(key or "").strip()
        }

    positive_types = {source_type for source_type, count in source_type_counts.items() if count > 0}
    if positive_types and positive_types <= {"kb"}:
        return COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES
    return COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS


def get_collection_readiness_status(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
) -> CollectionReadinessStatus:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    documents: list[Any] = []
    summary: dict[str, Any] = {}
    try:
        documents = _list_collection_documents(
            stores,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
        summary = _collection_document_summary(
            stores,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            records=documents,
        )
    except Exception as exc:
        return CollectionReadinessStatus(
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            maintenance_policy=COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS,
            document_count=0,
            ready=False,
            reason="collection_status_check_failed",
            sync_error=str(exc),
            suggested_fix="",
        )

    collection_store = getattr(stores, "collection_store", None)
    record = None
    if collection_store is not None and hasattr(collection_store, "get_collection"):
        try:
            record = collection_store.get_collection(effective_collection_id, tenant_id=tenant_id)
        except Exception:
            record = None
    maintenance_policy = infer_collection_maintenance_policy(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        collection_record=record,
        collection_summary=summary,
        documents=documents,
    )
    document_count = int(summary.get("document_count") or len(documents))
    if maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
        coverage = get_kb_coverage_status(
            settings,
            stores,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
        return CollectionReadinessStatus(
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            maintenance_policy=maintenance_policy,
            document_count=document_count,
            ready=bool(coverage.ready),
            reason=str(coverage.reason or "ready"),
            configured_source_paths=tuple(coverage.configured_source_paths),
            missing_source_paths=tuple(coverage.missing_source_paths),
            indexed_source_paths=tuple(coverage.indexed_source_paths),
            sync_error=str(coverage.sync_error or ""),
            suggested_fix=str(coverage.suggested_fix or ""),
        )

    return CollectionReadinessStatus(
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        maintenance_policy=maintenance_policy,
        document_count=document_count,
        ready=document_count > 0,
        reason="indexed" if document_count > 0 else "empty_collection",
        suggested_fix=(
            ""
            if document_count > 0
            else f"Use the control panel upload or ingest actions to add documents to collection '{effective_collection_id}'."
        ),
    )


def _configured_paths_by_title(settings: Settings) -> dict[str, list[Path]]:
    grouped: dict[str, list[Path]] = {}
    if not hasattr(settings, "kb_dir"):
        return grouped
    for path in iter_kb_source_paths(settings):
        grouped.setdefault(path.name.casefold(), []).append(path.resolve())
    return grouped


def _configured_alias_path(
    title: str,
    configured_paths_by_title: dict[str, list[Path]],
) -> str:
    candidates = configured_paths_by_title.get(_normalize_title(title), [])
    if len(candidates) == 1:
        return str(candidates[0])
    return ""


def resolve_record_source_identity(
    record: Any,
    *,
    configured_paths_by_title: dict[str, list[Path]] | None = None,
) -> str:
    title = _record_title(record)
    source_type = _record_source_type(record).strip().casefold()
    if configured_paths_by_title and source_type == "kb":
        alias_path = _configured_alias_path(title, configured_paths_by_title)
        if alias_path:
            return f"path:{_normalize_path_text(alias_path)}"
    source_identity = _normalize_source_identity_text(_record_source_identity(record))
    if source_identity:
        return source_identity
    source_path = _record_source_path(record)
    if source_path:
        return f"path:{_normalize_path_text(source_path)}"
    if title:
        return f"title:{_normalize_title(title)}"
    return "unknown:"


def resolve_candidate_source_identity(
    *,
    path: Path,
    title: str,
    configured_paths_by_title: dict[str, list[Path]] | None = None,
    source_identity: str = "",
) -> str:
    if source_identity:
        normalized_identity = _normalize_source_identity_text(str(source_identity))
        if normalized_identity:
            return normalized_identity
    if configured_paths_by_title:
        alias_path = _configured_alias_path(title, configured_paths_by_title)
        if alias_path:
            return f"path:{_normalize_path_text(alias_path)}"
    return f"path:{_normalize_path_text(str(path))}"


def _parse_ingested_at(value: str) -> dt.datetime:
    text = str(value or "").strip()
    if not text:
        return dt.datetime.min.replace(tzinfo=dt.timezone.utc)
    try:
        normalized = text.replace("Z", "+00:00")
        parsed = dt.datetime.fromisoformat(normalized)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=dt.timezone.utc)
        return parsed
    except ValueError:
        return dt.datetime.min.replace(tzinfo=dt.timezone.utc)


def select_active_kb_record(records: Sequence[Any]) -> Any | None:
    if not records:
        return None
    explicit_active = [record for record in records if _record_active(record)]
    if explicit_active:
        records = explicit_active
    return sorted(
        records,
        key=lambda item: (
            _record_version_ordinal(item),
            _parse_ingested_at(str(getattr(item, "ingested_at", "") or "")),
            str(getattr(item, "doc_id", "") or ""),
        ),
        reverse=True,
    )[0]


def _record_active(record: Any) -> bool:
    if isinstance(record, dict):
        value = record.get("active", True)
    else:
        value = getattr(record, "active", True)
    return bool(True if value is None else value)


def _record_version_ordinal(record: Any) -> int:
    if isinstance(record, dict):
        value = record.get("version_ordinal", 1)
    else:
        value = getattr(record, "version_ordinal", 1)
    try:
        return max(1, int(value or 1))
    except (TypeError, ValueError):
        return 1


def _record_metadata_confidence(record: Any) -> float:
    value = record.get("metadata_confidence") if isinstance(record, dict) else getattr(record, "metadata_confidence", None)
    if value is None:
        source_metadata = record.get("source_metadata", {}) if isinstance(record, dict) else getattr(record, "source_metadata", {})
        if isinstance(source_metadata, dict):
            value = source_metadata.get("metadata_confidence")
    try:
        return max(0.0, min(1.0, float(value if value is not None else 0.5)))
    except (TypeError, ValueError):
        return 0.5


def _record_parser_chain(record: Any) -> list[str]:
    provenance = record.get("parser_provenance", {}) if isinstance(record, dict) else getattr(record, "parser_provenance", {})
    source_metadata = record.get("source_metadata", {}) if isinstance(record, dict) else getattr(record, "source_metadata", {})
    if not isinstance(provenance, dict) and isinstance(source_metadata, dict):
        provenance = source_metadata.get("parser_provenance", {})
    if isinstance(provenance, dict):
        chain = [str(item) for item in list(provenance.get("chain") or []) if str(item)]
        if chain:
            return chain
    if isinstance(source_metadata, dict):
        index_metadata = source_metadata.get("index_metadata")
        if isinstance(index_metadata, dict):
            return [str(item) for item in list(index_metadata.get("parser_chain") or []) if str(item)]
    return []


def _record_parser_warnings(records: Sequence[Any]) -> list[str]:
    warnings: list[str] = []
    for record in records:
        source_metadata = record.get("source_metadata", {}) if isinstance(record, dict) else getattr(record, "source_metadata", {})
        if not isinstance(source_metadata, dict):
            continue
        for warning in list(source_metadata.get("parser_warnings") or []):
            text = str(warning or "").strip()
            if text and text not in warnings:
                warnings.append(text)
    return warnings


def _confidence_bucket(value: float) -> str:
    if value >= 0.8:
        return "high"
    if value >= 0.55:
        return "medium"
    return "low"


def _kb_document_version(record: Any, *, active: bool) -> KBDocumentVersion:
    source_metadata = dict(getattr(record, "source_metadata", {}) or {})
    provenance = getattr(record, "parser_provenance", {}) or source_metadata.get("parser_provenance")
    parser_chain: list[str] = []
    if isinstance(provenance, dict):
        parser_chain = [str(item) for item in list(provenance.get("chain") or []) if str(item)]
    if not parser_chain:
        parser_chain = _record_parser_chain(record)
    signal_summary = getattr(record, "signal_summary", {}) or source_metadata.get("signal_summary") or {}
    return KBDocumentVersion(
        doc_id=str(getattr(record, "doc_id", "") or ""),
        title=_record_title(record),
        source_type=str(getattr(record, "source_type", "") or ""),
        source_path=str(getattr(record, "source_path", "") or ""),
        collection_id=str(getattr(record, "collection_id", "") or ""),
        content_hash=str(getattr(record, "content_hash", "") or ""),
        ingested_at=str(getattr(record, "ingested_at", "") or ""),
        num_chunks=int(getattr(record, "num_chunks", 0) or 0),
        file_type=str(getattr(record, "file_type", "") or ""),
        doc_structure_type=str(getattr(record, "doc_structure_type", "") or ""),
        active=active,
        version_ordinal=_record_version_ordinal(record),
        superseded_at=str(getattr(record, "superseded_at", "") or ""),
        parser_chain=tuple(parser_chain),
        extraction_status=str(getattr(record, "extraction_status", "") or source_metadata.get("extraction_status") or "success"),
        extraction_error=str(getattr(record, "extraction_error", "") or source_metadata.get("extraction_error") or ""),
        metadata_confidence=_record_metadata_confidence(record),
        lifecycle_phase=str(getattr(record, "lifecycle_phase", "") or source_metadata.get("lifecycle_phase") or ""),
        doc_type=str(getattr(record, "doc_type", "") or source_metadata.get("doc_type") or ""),
        program_entities=tuple(str(item) for item in list(getattr(record, "program_entities", []) or source_metadata.get("program_entities") or []) if str(item)),
        signal_summary=dict(signal_summary) if isinstance(signal_summary, dict) else {},
    )


def _normalized_source_paths(paths: Sequence[Path]) -> tuple[str, ...]:
    return tuple(sorted(str(path.resolve()) for path in paths))


def _record_title(record: Any) -> str:
    if isinstance(record, dict):
        return str(record.get("title") or "")
    return str(getattr(record, "title", "") or "")


def _record_source_type(record: Any) -> str:
    if isinstance(record, dict):
        return str(record.get("source_type") or "")
    return str(getattr(record, "source_type", "") or "")


def _record_source_path(record: Any) -> str:
    if isinstance(record, dict):
        value = record.get("source_path")
    else:
        value = getattr(record, "source_path", "")
    if not value:
        return ""
    try:
        return str(Path(str(value)).resolve())
    except Exception:
        return str(value)


def _record_has_source_reference(record: Any, *, configured_source_path: str = "") -> bool:
    if configured_source_path and Path(configured_source_path).exists():
        return True
    source_path = str(record.get("source_path") if isinstance(record, dict) else getattr(record, "source_path", "") or "")
    source_uri = str(record.get("source_uri") if isinstance(record, dict) else getattr(record, "source_uri", "") or "")
    storage_backend = str(
        record.get("source_storage_backend") if isinstance(record, dict) else getattr(record, "source_storage_backend", "")
    ).strip().lower()
    object_key = str(record.get("source_object_key") if isinstance(record, dict) else getattr(record, "source_object_key", "") or "")
    object_bucket = str(record.get("source_object_bucket") if isinstance(record, dict) else getattr(record, "source_object_bucket", "") or "")
    source_metadata = record.get("source_metadata", {}) if isinstance(record, dict) else getattr(record, "source_metadata", {})
    blob_ref = source_metadata.get("blob_ref") if isinstance(source_metadata, dict) and isinstance(source_metadata.get("blob_ref"), dict) else {}
    if storage_backend and storage_backend != "local":
        return bool(source_uri or object_key or blob_ref.get("uri") or blob_ref.get("key"))
    if "://" in source_path or "://" in source_uri:
        return True
    return bool(source_path and Path(source_path).exists())


def _record_source_identity(record: Any) -> str:
    if isinstance(record, dict):
        value = record.get("source_identity")
    else:
        value = getattr(record, "source_identity", "")
    return str(value or "").strip()


def _resolve_path_overrides(mapping: dict[str, str] | None) -> dict[str, str]:
    resolved: dict[str, str] = {}
    for raw_path, override in (mapping or {}).items():
        try:
            key = str(Path(raw_path).expanduser().resolve())
        except Exception:
            key = str(raw_path)
        resolved[key] = str(override or "")
    return resolved


def build_kb_coverage_status(
    settings: Settings,
    indexed_records: Sequence[Any],
    *,
    tenant_id: str,
    collection_id: str | None = None,
    sync_attempted: bool = False,
    sync_error: str = "",
    synced_doc_ids: Sequence[str] = (),
) -> KBCoverageStatus:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    configured_paths = iter_kb_source_paths(settings)
    configured_source_paths = _normalized_source_paths(configured_paths)
    indexed_source_paths = tuple(
        sorted(
            {
                source_path
                for source_path in (_record_source_path(record) for record in indexed_records)
                if source_path
            }
        )
    )
    indexed_titles = {
        title
        for title in (_record_title(record) for record in indexed_records)
        if title
    }
    missing_source_paths = tuple(
        path
        for path in configured_source_paths
        if path not in indexed_source_paths and Path(path).name not in indexed_titles
    )
    return KBCoverageStatus(
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        configured_source_paths=configured_source_paths,
        missing_source_paths=missing_source_paths,
        indexed_source_paths=indexed_source_paths,
        indexed_doc_count=len(indexed_records),
        sync_attempted=sync_attempted,
        sync_error=str(sync_error or "").strip(),
        synced_doc_ids=tuple(str(doc_id) for doc_id in synced_doc_ids if str(doc_id)),
    )


def get_kb_coverage_status(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
) -> KBCoverageStatus:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    list_documents = getattr(getattr(stores, "doc_store", None), "list_documents", None)
    if not callable(list_documents):
        return build_kb_coverage_status(
            settings,
            [],
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
    try:
        records = list_documents(
            source_type="kb",
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
    except Exception as exc:
        return build_kb_coverage_status(
            settings,
            [],
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            sync_error=str(exc),
        )
    return build_kb_coverage_status(
        settings,
        list(records),
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
    )


def build_collection_health_report(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
    source_type: str = "",
    maintenance_policy: str = "",
) -> KBCorpusHealthReport:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    source_type_filter = str(source_type or "").strip()
    resolved_maintenance_policy = normalize_collection_maintenance_policy(
        str(maintenance_policy or "").strip(),
        default="",
    )
    try:
        records = _list_documents_for_ingest(
            stores,
            source_type=source_type_filter,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            active_only=False,
        )
    except Exception as exc:
        if not resolved_maintenance_policy:
            resolved_maintenance_policy = infer_collection_maintenance_policy(
                settings,
                stores,
                tenant_id=tenant_id,
                collection_id=effective_collection_id,
                documents=[],
            )
        configured_source_paths: tuple[str, ...] = ()
        missing_source_paths: tuple[str, ...] = ()
        if resolved_maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
            coverage = build_kb_coverage_status(
                settings,
                [],
                tenant_id=tenant_id,
                collection_id=effective_collection_id,
                sync_error=str(exc),
            )
            configured_source_paths = coverage.configured_source_paths
            missing_source_paths = coverage.missing_source_paths
        return KBCorpusHealthReport(
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            configured_source_paths=configured_source_paths,
            missing_source_paths=missing_source_paths,
            indexed_doc_count=0,
            active_doc_count=0,
            maintenance_policy=resolved_maintenance_policy or COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS,
            sync_error=str(exc),
        )

    if not resolved_maintenance_policy:
        resolved_maintenance_policy = infer_collection_maintenance_policy(
            settings,
            stores,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            documents=records,
        )
    configured_source_paths: tuple[str, ...] = ()
    missing_source_paths: tuple[str, ...] = ()
    sync_error = ""
    configured_paths_by_title: dict[str, list[Path]] = {}
    active_records = [record for record in records if _record_active(record)]
    if resolved_maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
        coverage = build_kb_coverage_status(
            settings,
            active_records,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
        configured_source_paths = coverage.configured_source_paths
        missing_source_paths = coverage.missing_source_paths
        sync_error = coverage.sync_error
        configured_paths_by_title = _configured_paths_by_title(settings)

    grouped: dict[str, list[Any]] = {}
    for record in records:
        source_identity = resolve_record_source_identity(
            record,
            configured_paths_by_title=configured_paths_by_title,
        )
        grouped.setdefault(source_identity, []).append(record)

    source_groups: list[KBSourceHealthGroup] = []
    duplicate_groups: list[KBSourceHealthGroup] = []
    drifted_groups: list[KBSourceHealthGroup] = []
    parser_counts: Counter[str] = Counter()
    doc_type_counts: Counter[str] = Counter()
    lifecycle_counts: Counter[str] = Counter()
    confidence_distribution: Counter[str] = Counter()
    stale_version_count = 0
    extraction_failure_count = 0
    low_confidence_metadata_count = 0
    missing_source_doc_count = 0
    parser_warning_count = 0

    for source_identity, group_records in sorted(grouped.items(), key=lambda item: item[0]):
        active = select_active_kb_record(group_records)
        if active is None:
            continue
        configured_source_path = _configured_alias_path(_record_title(active), configured_paths_by_title)
        candidate_path = configured_source_path or _record_source_path(active)
        source_exists = _record_has_source_reference(active, configured_source_path=str(configured_source_path or ""))
        current_file_hash = ""
        if candidate_path and Path(candidate_path).exists():
            try:
                current_file_hash = _file_hash(Path(candidate_path))
            except OSError:
                current_file_hash = ""
        content_drift = bool(
            current_file_hash
            and str(getattr(active, "content_hash", "") or "")
            and current_file_hash != str(getattr(active, "content_hash", "") or "")
        )
        sorted_group = sorted(
            group_records,
            key=lambda item: (
                _record_version_ordinal(item),
                _parse_ingested_at(str(getattr(item, "ingested_at", "") or "")),
                str(getattr(item, "doc_id", "") or ""),
            ),
            reverse=True,
        )
        active_group_records = [record for record in sorted_group if _record_active(record)]
        stale_records = [record for record in sorted_group if not _record_active(record)]
        duplicate_doc_ids = tuple(
            str(getattr(record, "doc_id", "") or "")
            for record in active_group_records
            if str(getattr(record, "doc_id", "") or "") and str(getattr(record, "doc_id", "") or "") != str(getattr(active, "doc_id", "") or "")
        )
        stale_doc_ids = tuple(str(getattr(record, "doc_id", "") or "") for record in stale_records if str(getattr(record, "doc_id", "") or ""))
        extraction_failure_doc_ids = tuple(
            str(getattr(record, "doc_id", "") or "")
            for record in active_group_records
            if str(getattr(record, "doc_id", "") or "")
            and str(getattr(record, "extraction_status", "") or "success").lower() not in {"success", "completed", "ok"}
        )
        low_confidence_doc_ids = tuple(
            str(getattr(record, "doc_id", "") or "")
            for record in active_group_records
            if str(getattr(record, "doc_id", "") or "") and _record_metadata_confidence(record) < 0.5
        )
        missing_source_doc_ids = tuple(
            str(getattr(record, "doc_id", "") or "")
            for record in active_group_records
            if str(getattr(record, "doc_id", "") or "")
            and not _record_has_source_reference(record, configured_source_path=str(configured_source_path or ""))
        )
        parser_warnings = tuple(_record_parser_warnings(active_group_records))
        for record in active_group_records:
            for parser_name in _record_parser_chain(record):
                parser_counts[parser_name] += 1
            doc_type_counts[str(getattr(record, "doc_type", "") or "unknown")] += 1
            lifecycle_counts[str(getattr(record, "lifecycle_phase", "") or "unknown")] += 1
            confidence_distribution[_confidence_bucket(_record_metadata_confidence(record))] += 1
        stale_version_count += len(stale_doc_ids)
        extraction_failure_count += len(extraction_failure_doc_ids)
        low_confidence_metadata_count += len(low_confidence_doc_ids)
        missing_source_doc_count += len(missing_source_doc_ids)
        parser_warning_count += len(parser_warnings)
        versions = tuple(
            _kb_document_version(record, active=_record_active(record))
            for record in sorted_group
        )
        group = KBSourceHealthGroup(
            source_identity=source_identity,
            title=_record_title(active),
            source_type=str(getattr(active, "source_type", "") or "kb"),
            collection_id=str(getattr(active, "collection_id", "") or effective_collection_id),
            configured_source_path=str(configured_source_path or ""),
            active_doc_id=str(getattr(active, "doc_id", "") or ""),
            active_content_hash=str(getattr(active, "content_hash", "") or ""),
            active_ingested_at=str(getattr(active, "ingested_at", "") or ""),
            active_source_path=str(getattr(active, "source_path", "") or ""),
            current_file_hash=current_file_hash,
            source_exists=source_exists,
            content_drift=content_drift,
            duplicate_doc_ids=duplicate_doc_ids,
            stale_version_doc_ids=stale_doc_ids,
            extraction_failure_doc_ids=extraction_failure_doc_ids,
            low_confidence_doc_ids=low_confidence_doc_ids,
            missing_source_doc_ids=missing_source_doc_ids,
            parser_warnings=parser_warnings,
            records=versions,
        )
        source_groups.append(group)
        if group.duplicate_doc_ids:
            duplicate_groups.append(group)
        if group.content_drift:
            drifted_groups.append(group)

    return KBCorpusHealthReport(
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        configured_source_paths=configured_source_paths,
        missing_source_paths=missing_source_paths,
        indexed_doc_count=len(records),
        active_doc_count=len(active_records),
        maintenance_policy=resolved_maintenance_policy or COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS,
        duplicate_groups=tuple(duplicate_groups),
        drifted_groups=tuple(drifted_groups),
        source_groups=tuple(source_groups),
        sync_error=sync_error,
        stale_version_count=stale_version_count,
        extraction_failure_count=extraction_failure_count,
        low_confidence_metadata_count=low_confidence_metadata_count,
        missing_source_doc_count=missing_source_doc_count,
        parser_warning_count=parser_warning_count,
        parser_counts=dict(parser_counts),
        doc_type_counts=dict(doc_type_counts),
        lifecycle_counts=dict(lifecycle_counts),
        metadata_confidence_distribution=dict(confidence_distribution),
    )


def build_kb_health_report(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
) -> KBCorpusHealthReport:
    return build_collection_health_report(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=collection_id,
        source_type="kb",
        maintenance_policy=COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES,
    )


def repair_collection_documents(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
    source_type: str = "",
    maintenance_policy: str = "",
    title_hint: str = "",
) -> KBRepairResult:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    health_before = build_collection_health_report(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        source_type=source_type,
        maintenance_policy=maintenance_policy,
    )
    title_filter = str(title_hint or "").strip().casefold()

    deleted_doc_ids: list[str] = []
    reindexed_doc_ids: list[str] = []
    ingested_missing_doc_ids: list[str] = []
    unresolved_paths: list[str] = []

    groups = list(health_before.source_groups)
    if title_filter:
        groups = [group for group in groups if title_filter in group.title.casefold()]

    for group in groups:
        for doc_id in group.duplicate_doc_ids:
            stores.doc_store.delete_document(doc_id, tenant_id=tenant_id)
            deleted_doc_ids.append(doc_id)

    for group in groups:
        if not group.content_drift:
            continue
        candidate_path = str(group.configured_source_path or group.active_source_path or "")
        if not candidate_path or not Path(candidate_path).exists():
            unresolved_paths.append(candidate_path or group.title)
            continue
        doc_ids = ingest_paths(
            settings,
            stores,
            [Path(candidate_path)],
            source_type=group.source_type or source_type or "kb",
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
        reindexed_doc_ids.extend(doc_ids)

    if health_before.maintenance_policy == COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES:
        missing_paths = [
            Path(path)
            for path in health_before.missing_source_paths
            if not title_filter or title_filter in Path(path).name.casefold()
        ]
        if missing_paths:
            ingested_missing_doc_ids.extend(
                ingest_paths(
                    settings,
                    stores,
                    missing_paths,
                    source_type=source_type or "kb",
                    tenant_id=tenant_id,
                    collection_id=effective_collection_id,
                )
            )

    health_after = build_collection_health_report(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        source_type=source_type,
        maintenance_policy=health_before.maintenance_policy,
    )
    return KBRepairResult(
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        deleted_doc_ids=tuple(dict.fromkeys(item for item in deleted_doc_ids if item)),
        reindexed_doc_ids=tuple(dict.fromkeys(item for item in reindexed_doc_ids if item)),
        ingested_missing_doc_ids=tuple(dict.fromkeys(item for item in ingested_missing_doc_ids if item)),
        unresolved_paths=tuple(dict.fromkeys(item for item in unresolved_paths if item)),
        health_before=health_before,
        health_after=health_after,
    )


def repair_kb_collection(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
    title_hint: str = "",
) -> KBRepairResult:
    return repair_collection_documents(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=collection_id,
        source_type="kb",
        maintenance_policy=COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES,
        title_hint=title_hint,
    )


def _file_hash(path: Path) -> str:
    sha = hashlib.sha1()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            sha.update(chunk)
    return sha.hexdigest()


def _load_documents_with_docling(path: Path) -> List[Document]:
    try:
        from docling.document_converter import DocumentConverter
    except Exception as exc:
        logger.debug("Docling is unavailable for %s: %s", path, exc)
        return []

    try:
        converter = DocumentConverter()
        result = converter.convert(str(path))
        document = getattr(result, "document", None)
        if document is None:
            return []

        rendered_text = ""
        for method_name in ("export_to_markdown", "export_to_text"):
            method = getattr(document, method_name, None)
            if not callable(method):
                continue
            try:
                content = method()
            except TypeError:
                continue
            if content:
                rendered_text = str(content)
                break

        if not rendered_text:
            rendered_text = str(document or "")
        if not rendered_text.strip():
            return []

        return [Document(page_content=rendered_text, metadata={"parser": "docling"})]
    except Exception as exc:
        logger.warning("Docling extraction failed for %s: %s", path, exc)
        return []


def _load_docx_with_python_docx(path: Path) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n\n".join(parts).strip()
    if not content:
        return []
    return [Document(page_content=content, metadata={"parser": "python-docx"})]


def _load_document_bundle(
    path: Path,
    settings: Settings,
    *,
    parser_strategy: str = "docling_primary",
) -> ParsedDocumentBundle:
    return load_documents_with_parsers(path, settings, parser_strategy=parser_strategy)


def _load_documents(path: Path, settings: Settings) -> List[Document]:
    return _load_document_bundle(path, settings).documents


def _chunk_size(settings: Settings) -> int:
    return int(getattr(settings, "chunk_size", 900) or 900)


def _chunk_overlap(settings: Settings) -> int:
    return int(getattr(settings, "chunk_overlap", 150) or 150)


def _general_split(settings: Settings, docs: List[Document]) -> List[Document]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=_chunk_size(settings),
        chunk_overlap=_chunk_overlap(settings),
        add_start_index=True,
    )
    return splitter.split_documents(docs)


def _split_with_structure(
    settings: Settings,
    docs: List[Document],
    structure: StructureAnalysis,
) -> List[Document]:
    if any(bool((doc.metadata or {}).get("is_prechunked")) for doc in docs):
        return list(docs)

    if structure.has_clauses:
        chunks: List[Document] = []
        for doc in docs:
            chunks.extend(
                clause_split(
                    doc,
                    max_clause_chars=_chunk_size(settings) * 2,
                    overlap_chars=_chunk_overlap(settings),
                )
            )
    else:
        chunks = _general_split(settings, docs)
        for chunk in chunks:
            if PROCESS_FLOW_PATTERN.search(chunk.page_content):
                chunk.metadata["chunk_type"] = "process_flow"
            elif REQUIREMENT_PATTERN.search(chunk.page_content):
                chunk.metadata["chunk_type"] = "requirement"
            elif not chunk.metadata.get("chunk_type"):
                chunk.metadata["chunk_type"] = "general"

    return chunks


def _build_chunk_records(
    chunks: List[Document],
    doc_id: str,
    *,
    collection_id: str,
    document_index_metadata: DocumentIndexMetadata | dict[str, Any] | None = None,
) -> List[ChunkRecord]:
    records: List[ChunkRecord] = []
    seen_indices: set[int] = set()
    for index, chunk in enumerate(chunks):
        metadata = chunk.metadata or {}
        try:
            chunk_index = int(metadata.get("chunk_index", index))
        except (TypeError, ValueError):
            chunk_index = index
        if chunk_index in seen_indices:
            chunk_index = index
            while chunk_index in seen_indices:
                chunk_index += 1
        seen_indices.add(chunk_index)
        chunk_id = f"{doc_id}#chunk{chunk_index:04d}"
        chunk_metadata = (
            build_chunk_index_metadata(
                chunk,
                chunk_id=chunk_id,
                chunk_index=chunk_index,
                document_metadata=document_index_metadata,
            ).to_dict()
            if document_index_metadata is not None
            else {}
        )
        status_metadata = metadata.get("status_workbook") if isinstance(metadata.get("status_workbook"), dict) else {}
        status_domains = [str(item) for item in (metadata.get("status_domains") or []) if str(item)]
        if status_metadata:
            chunk_metadata["status_workbook"] = dict(status_metadata)
        if status_domains:
            chunk_metadata["status_domains"] = status_domains
        records.append(
            ChunkRecord(
                chunk_id=chunk_id,
                doc_id=doc_id,
                collection_id=collection_id,
                chunk_index=chunk_index,
                content=chunk.page_content,
                chunk_type=str(metadata.get("chunk_type", "general")),
                page_number=metadata.get("page"),
                clause_number=metadata.get("clause_number") or None,
                section_title=metadata.get("section_title") or None,
                sheet_name=metadata.get("sheet_name") or None,
                row_start=metadata.get("row_start"),
                row_end=metadata.get("row_end"),
                cell_range=metadata.get("cell_range") or None,
                metadata_json=chunk_metadata,
                embedding_text=_metadata_embedding_text(chunk.page_content, chunk_metadata),
                embedding=None,
            )
        )
    return records


def _metadata_embedding_text(content: str, metadata: dict[str, Any]) -> str:
    if not metadata or str(metadata.get("metadata_profile") or "") == "off":
        return content
    location = dict(metadata.get("location") or {})
    parts: list[str] = []
    for label, value in (
        ("type", metadata.get("document_type")),
        ("chunk", metadata.get("chunk_type")),
        ("tags", ", ".join(str(tag) for tag in list(metadata.get("tags") or [])[:6])),
        ("status_domains", ", ".join(str(tag) for tag in list(metadata.get("status_domains") or [])[:8])),
        ("section", " > ".join(str(item) for item in list(metadata.get("section_path") or [])[:4])),
        ("clause", location.get("clause_number")),
        ("page", location.get("page_number")),
        ("sheet", location.get("sheet_name")),
        ("cells", location.get("cell_range")),
    ):
        text = str(value or "").strip()
        if text:
            parts.append(f"{label}: {text}")
    if not parts:
        return content
    return f"[index metadata] {' | '.join(parts)}\n\n{content}"


def _workbook_profile_metadata(path: Path) -> dict[str, Any]:
    if path.suffix.lower() not in {".xlsx", ".xls"}:
        return {}
    try:
        return profile_workbook(path).to_dict()
    except Exception as exc:
        return {
            "workbook_name": path.name,
            "source_path": str(path),
            "file_type": path.suffix.lstrip(".").lower(),
            "warnings": [f"Workbook profile failed: {exc}"],
        }


def _document_record_from_value(record: Any) -> DocumentRecord:
    return DocumentRecord(
        doc_id=str(getattr(record, "doc_id", "") or ""),
        tenant_id=str(getattr(record, "tenant_id", "") or "local-dev"),
        collection_id=str(getattr(record, "collection_id", "") or "default"),
        title=str(getattr(record, "title", "") or ""),
        source_type=str(getattr(record, "source_type", "") or ""),
        content_hash=str(getattr(record, "content_hash", "") or ""),
        source_path=str(getattr(record, "source_path", "") or ""),
        num_chunks=int(getattr(record, "num_chunks", 0) or 0),
        ingested_at=str(getattr(record, "ingested_at", "") or ""),
        file_type=str(getattr(record, "file_type", "") or ""),
        doc_structure_type=str(getattr(record, "doc_structure_type", "") or "general"),
        source_display_path=str(getattr(record, "source_display_path", "") or ""),
        source_identity=str(getattr(record, "source_identity", "") or ""),
        source_metadata=dict(getattr(record, "source_metadata", {}) or {}),
        source_uri=str(getattr(record, "source_uri", "") or ""),
        source_storage_backend=str(getattr(record, "source_storage_backend", "") or ""),
        source_object_bucket=str(getattr(record, "source_object_bucket", "") or ""),
        source_object_key=str(getattr(record, "source_object_key", "") or ""),
        source_etag=str(getattr(record, "source_etag", "") or ""),
        source_size_bytes=int(getattr(record, "source_size_bytes", 0) or 0),
        source_content_type=str(getattr(record, "source_content_type", "") or ""),
        active=bool(True if getattr(record, "active", True) is None else getattr(record, "active", True)),
        version_ordinal=_record_version_ordinal(record),
        superseded_at=str(getattr(record, "superseded_at", "") or ""),
        parser_provenance=dict(getattr(record, "parser_provenance", {}) or {}),
        extraction_status=str(getattr(record, "extraction_status", "") or "success"),
        extraction_error=str(getattr(record, "extraction_error", "") or ""),
        metadata_confidence=_record_metadata_confidence(record),
        lifecycle_phase=str(getattr(record, "lifecycle_phase", "") or ""),
        doc_type=str(getattr(record, "doc_type", "") or ""),
        program_entities=[str(item) for item in list(getattr(record, "program_entities", []) or []) if str(item)],
        signal_summary=dict(getattr(record, "signal_summary", {}) or {}),
    )


def _sync_requirement_inventory_for_document(
    stores: KnowledgeStores,
    document: DocumentRecord,
    *,
    tenant_id: str,
    chunk_records: Sequence[ChunkRecord] | None = None,
) -> List[RequirementStatementRecord]:
    requirement_store = getattr(stores, "requirement_store", None)
    if requirement_store is None:
        return []
    if not supports_requirements_extraction(document.file_type):
        requirement_store.delete_doc_statements(document.doc_id, tenant_id)
        return []
    source_chunks = list(chunk_records) if chunk_records is not None else list(
        stores.chunk_store.list_document_chunks(document.doc_id, tenant_id)
    )
    statements = build_requirement_statement_records(document, source_chunks)
    requirement_store.replace_doc_statements(
        document.doc_id,
        tenant_id,
        statements=statements,
    )
    return statements


def backfill_requirement_statements(
    settings: Settings,
    stores: KnowledgeStores,
    *,
    tenant_id: str,
    collection_id: str | None = None,
    source_type: str = "",
    title_hint: str = "",
) -> RequirementBackfillResult:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    records = list(
        stores.doc_store.list_documents(
            tenant_id=tenant_id,
            source_type=str(source_type or ""),
            collection_id=effective_collection_id,
        )
    )
    if str(source_type or "").strip() == "kb":
        configured_paths_by_title = _configured_paths_by_title(settings)
        grouped: dict[str, list[Any]] = {}
        for record in records:
            identity = resolve_record_source_identity(
                record,
                configured_paths_by_title=configured_paths_by_title,
            )
            grouped.setdefault(identity, []).append(record)
        records = [record for record in (select_active_kb_record(items) for items in grouped.values()) if record is not None]

    title_filter = str(title_hint or "").strip().casefold()
    if title_filter:
        records = [
            record
            for record in records
            if title_filter in str(getattr(record, "title", "") or "").casefold()
        ]

    processed_doc_ids: list[str] = []
    statement_count = 0
    unsupported_documents: list[dict[str, str]] = []
    for record in records:
        document = _document_record_from_value(record)
        if not supports_requirements_extraction(document.file_type):
            unsupported_documents.append(
                {
                    "doc_id": document.doc_id,
                    "title": document.title,
                    "file_type": document.file_type,
                }
            )
            continue
        statements = _sync_requirement_inventory_for_document(
            stores,
            document,
            tenant_id=tenant_id,
        )
        processed_doc_ids.append(document.doc_id)
        statement_count += len(statements)

    return RequirementBackfillResult(
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        source_type=str(source_type or ""),
        processed_doc_ids=tuple(processed_doc_ids),
        statement_count=statement_count,
        unsupported_documents=tuple(unsupported_documents),
    )


def _expand_ingest_paths(paths: Iterable[Path]) -> List[Path]:
    resolved: List[Path] = []
    seen: set[Path] = set()
    for raw_path in paths:
        path = Path(raw_path).expanduser()
        if not path.exists():
            continue
        candidates = [path]
        if path.is_dir():
            candidates = [item for item in sorted(path.rglob("*")) if item.is_file()]
        for candidate in candidates:
            resolved_candidate = candidate.resolve()
            if resolved_candidate in seen:
                continue
            seen.add(resolved_candidate)
            resolved.append(resolved_candidate)
    return resolved


def _record_index_metadata_matches(record: Any, *, metadata_profile: str, metadata_enrichment: str = "") -> bool:
    profile = normalize_metadata_profile(metadata_profile)
    enrichment = str(metadata_enrichment or "").strip().lower()
    source_metadata = dict(getattr(record, "source_metadata", {}) or {})
    index_metadata = source_metadata.get("index_metadata")
    if not isinstance(index_metadata, dict):
        return False
    matches = (
        str(index_metadata.get("extractor_version") or "") == INDEX_METADATA_VERSION
        and str(index_metadata.get("metadata_profile") or "") == profile
    )
    if enrichment:
        stored_enrichment = str(source_metadata.get("metadata_enrichment") or index_metadata.get("metadata_enrichment") or "").strip().lower()
        matches = matches and stored_enrichment == enrichment
    return matches


def _list_documents_for_ingest(
    stores: KnowledgeStores,
    *,
    source_type: str = "",
    tenant_id: str,
    collection_id: str,
    active_only: bool = True,
) -> list[Any]:
    list_documents = getattr(getattr(stores, "doc_store", None), "list_documents", None)
    if not callable(list_documents):
        return []
    try:
        return list(
            list_documents(
                source_type=source_type,
                tenant_id=tenant_id,
                collection_id=collection_id,
                active_only=active_only,
            )
        )
    except TypeError:
        return list(
            list_documents(
                source_type=source_type,
                tenant_id=tenant_id,
                collection_id=collection_id,
            )
        )


def _supersede_prior_source_versions(
    stores: KnowledgeStores,
    records: list[Any],
    *,
    tenant_id: str,
    collection_id: str,
    source_type: str,
    source_identity: str,
    active_doc_id: str,
    superseded_at: str,
) -> None:
    supersede = getattr(getattr(stores, "doc_store", None), "supersede_source_versions", None)
    if callable(supersede):
        supersede(
            tenant_id=tenant_id,
            collection_id=collection_id,
            source_type=source_type,
            source_identity=source_identity,
            active_doc_id=active_doc_id,
            superseded_at=superseded_at,
        )
    for record in records:
        if str(getattr(record, "doc_id", "") or "") == active_doc_id:
            continue
        try:
            setattr(record, "active", False)
            setattr(record, "superseded_at", superseded_at)
        except Exception:
            continue


def preview_path_index_metadata(
    settings: Settings,
    path: Path,
    *,
    metadata_profile: str = "auto",
    metadata_enrichment: str = "llm",
    parser_strategy: str = "docling_primary",
    providers: object | None = None,
    source_metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Return deterministic indexing metadata for a path without writing records."""
    resolved_path = Path(path).expanduser().resolve()
    source_metadata = dict(source_metadata or {})
    workbook_profile = _workbook_profile_metadata(resolved_path)
    if workbook_profile:
        source_metadata["workbook_profile"] = workbook_profile
    bundle = _load_document_bundle(resolved_path, settings, parser_strategy=parser_strategy)
    raw_docs = bundle.documents
    if not raw_docs:
        return {
            "path": str(resolved_path),
            "title": resolved_path.name,
            "metadata_summary": summarize_index_metadata([], metadata_profile=metadata_profile).to_dict(),
            "document_metadata": {},
            "chunk_count": 0,
            "parser_provenance": bundle.parser_provenance,
            "extraction_status": bundle.extraction_status,
            "extraction_error": bundle.extraction_error,
            "warnings": ["No extractable content was found.", *bundle.warnings, *bundle.errors],
        }
    for doc in raw_docs:
        doc.metadata = {
            **(doc.metadata or {}),
            "title": resolved_path.name,
            "source_path": str(resolved_path),
            "file_type": resolved_path.suffix.lstrip(".").lower(),
        }
    full_text = " ".join(doc.page_content for doc in raw_docs)
    structure = detect_structure(full_text)
    document_metadata = build_document_index_metadata(
        path=resolved_path,
        raw_docs=raw_docs,
        structure=structure,
        metadata_profile=metadata_profile,
        metadata_enrichment=metadata_enrichment,
        source_metadata=source_metadata,
    )
    corpus_metadata = enrich_corpus_metadata(
        path=resolved_path,
        raw_docs=raw_docs,
        source_metadata={
            **dict(source_metadata or {}),
            **bundle.to_source_metadata(),
        },
        document_metadata=document_metadata.to_dict(),
        providers=providers,
        metadata_enrichment=metadata_enrichment,
    )
    chunks = _split_with_structure(settings, raw_docs, structure)
    chunk_metadata = [
        build_chunk_index_metadata(
            chunk,
            chunk_id=f"preview#chunk{index:04d}",
            chunk_index=index,
            document_metadata=document_metadata,
        )
        for index, chunk in enumerate(chunks[:500])
    ]
    return {
        "path": str(resolved_path),
        "title": resolved_path.name,
        "document_metadata": document_metadata.to_dict(),
        "corpus_metadata": corpus_metadata,
        "parser_provenance": bundle.parser_provenance,
        "extraction_status": bundle.extraction_status,
        "extraction_error": bundle.extraction_error,
        "metadata_summary": summarize_index_metadata(
            [document_metadata],
            chunk_metadata,
            metadata_profile=metadata_profile,
        ).to_dict(),
        "chunk_count": len(chunks),
        "preview_chunks": [
            {
                "chunk_index": index,
                "chunk_type": str((chunk.metadata or {}).get("chunk_type") or "general"),
                "section_title": str((chunk.metadata or {}).get("section_title") or ""),
                "text_preview": str(chunk.page_content or "")[:500],
            }
            for index, chunk in enumerate(chunks[:20])
        ],
        "warnings": [*list(document_metadata.warnings), *bundle.warnings, *list(corpus_metadata.get("warnings") or [])],
    }


def ingest_paths(
    settings: Settings,
    stores: KnowledgeStores,
    paths: Iterable[Path],
    *,
    source_type: str,
    tenant_id: str,
    collection_id: str | None = None,
    source_display_paths: dict[str, str] | None = None,
    source_identities: dict[str, str] | None = None,
    source_metadata_by_path: dict[str, dict[str, Any]] | None = None,
    metadata_profile: str = "auto",
    metadata_enrichment: str = "llm",
    parser_strategy: str = "docling_primary",
    providers: object | None = None,
) -> List[str]:
    ingested_doc_ids: List[str] = []
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    normalized_metadata_profile = normalize_metadata_profile(metadata_profile)
    normalized_metadata_enrichment = str(metadata_enrichment or "deterministic").strip().lower() or "deterministic"
    if getattr(stores, "collection_store", None) is not None:
        stores.collection_store.ensure_collection(
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            maintenance_policy=(
                COLLECTION_MAINTENANCE_CONFIGURED_KB_SOURCES
                if str(source_type or "").strip().lower() == "kb"
                else COLLECTION_MAINTENANCE_INDEXED_DOCUMENTS
            ),
        )
    existing_records = _list_documents_for_ingest(
        stores,
        source_type=source_type,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
        active_only=False,
    )
    configured_paths_by_title = _configured_paths_by_title(settings) if source_type == "kb" else {}
    resolved_display_paths = _resolve_path_overrides(source_display_paths)
    resolved_source_identities = _resolve_path_overrides(source_identities)
    resolved_source_metadata = {
        str(Path(key).expanduser().resolve()): dict(value or {})
        for key, value in dict(source_metadata_by_path or {}).items()
        if str(key).strip() and isinstance(value, dict)
    }
    for path in _expand_ingest_paths(paths):
        file_hash = _file_hash(path)
        title = path.name
        source_key = str(path.resolve())
        source_display_path = resolved_display_paths.get(source_key) or title
        provided_source_identity = resolved_source_identities.get(source_key) or ""
        candidate_identity = resolve_candidate_source_identity(
            path=path,
            title=title,
            configured_paths_by_title=configured_paths_by_title,
            source_identity=provided_source_identity,
        )
        source_metadata = {
            **dict(resolved_source_metadata.get(source_key) or {}),
            "source_display_path": source_display_path,
            "source_identity": candidate_identity,
        }
        workbook_profile = _workbook_profile_metadata(path)
        if workbook_profile:
            source_metadata["workbook_profile"] = workbook_profile
        blob_ref = source_metadata.get("blob_ref") if isinstance(source_metadata.get("blob_ref"), dict) else {}
        source_uri = str(source_metadata.get("source_uri") or blob_ref.get("uri") or "").strip()
        source_storage_backend = str(blob_ref.get("backend") or "").strip()
        if not source_storage_backend:
            source_storage_backend = "local"
        stored_source_path = str(path)
        if source_uri and source_storage_backend != "local":
            stored_source_path = source_uri
        elif source_storage_backend == "local" and blob_ref.get("key"):
            stored_source_path = str(blob_ref.get("key") or path)
        same_source_records = [
            record
            for record in existing_records
            if resolve_record_source_identity(
                record,
                configured_paths_by_title=configured_paths_by_title,
            ) == candidate_identity
        ]
        active_existing = select_active_kb_record(same_source_records)
        if active_existing is not None and str(getattr(active_existing, "content_hash", "") or "") == file_hash:
            if _record_index_metadata_matches(
                active_existing,
                metadata_profile=normalized_metadata_profile,
                metadata_enrichment=normalized_metadata_enrichment,
            ):
                _sync_requirement_inventory_for_document(
                    stores,
                    _document_record_from_value(active_existing),
                    tenant_id=tenant_id,
                )
                continue
        doc_id = make_doc_id(
            source_type=source_type,
            source_identity=candidate_identity,
            content_hash=file_hash,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )

        bundle = _load_document_bundle(path, settings, parser_strategy=parser_strategy)
        raw_docs = bundle.documents

        full_text = " ".join(doc.page_content for doc in raw_docs)
        structure = detect_structure(full_text)
        document_index_metadata = build_document_index_metadata(
            path=path,
            raw_docs=raw_docs,
            structure=structure,
            metadata_profile=normalized_metadata_profile,
            metadata_enrichment=normalized_metadata_enrichment,
            source_metadata=source_metadata,
        )
        corpus_metadata = enrich_corpus_metadata(
            path=path,
            raw_docs=raw_docs,
            source_metadata={
                **source_metadata,
                **bundle.to_source_metadata(),
            },
            document_metadata=document_index_metadata.to_dict(),
            providers=providers,
            metadata_enrichment=normalized_metadata_enrichment,
        )
        signal_summary = dict(corpus_metadata.get("signal_summary") or {})
        source_metadata = {
            **source_metadata,
            **bundle.to_source_metadata(),
            "metadata_profile": document_index_metadata.metadata_profile,
            "metadata_enrichment": corpus_metadata.get("metadata_enrichment") or document_index_metadata.metadata_enrichment,
            "metadata_extractor_version": document_index_metadata.extractor_version,
            "index_metadata": document_index_metadata.to_dict(),
            "parser_strategy": parser_strategy,
            "parser_provenance": bundle.parser_provenance,
            "corpus_metadata": corpus_metadata,
            "metadata_confidence": corpus_metadata.get("metadata_confidence", 0.5),
            "lifecycle_phase": corpus_metadata.get("lifecycle_phase") or "",
            "doc_type": corpus_metadata.get("doc_type") or "",
            "program_entities": list(corpus_metadata.get("program_entities") or []),
            "signal_summary": signal_summary,
        }
        for doc in raw_docs:
            doc.metadata = {
                **(doc.metadata or {}),
                "doc_id": doc_id,
                "title": title,
                "source_type": source_type,
                "source_path": stored_source_path,
                "source_display_path": source_display_path,
                "source_identity": candidate_identity,
                "source_metadata": source_metadata,
                "collection_id": effective_collection_id,
                "file_type": path.suffix.lstrip(".").lower(),
            }

        chunks = _split_with_structure(settings, raw_docs, structure)
        chunk_records = _build_chunk_records(
            chunks,
            doc_id,
            collection_id=effective_collection_id,
            document_index_metadata=document_index_metadata,
        )
        ingested_at = dt.datetime.now(dt.timezone.utc).isoformat()
        version_ordinal = max([_record_version_ordinal(record) for record in same_source_records] or [0]) + 1
        document_record = DocumentRecord(
            doc_id=doc_id,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            title=title,
            source_type=source_type,
            content_hash=file_hash,
            source_path=stored_source_path,
            num_chunks=len(chunk_records),
            ingested_at=ingested_at,
            file_type=path.suffix.lstrip(".").lower(),
            doc_structure_type=document_index_metadata.doc_structure_type,
            source_display_path=source_display_path,
            source_identity=candidate_identity,
            source_metadata=source_metadata,
            source_uri=source_uri or stored_source_path,
            source_storage_backend=source_storage_backend,
            source_object_bucket=str(blob_ref.get("bucket") or blob_ref.get("container") or ""),
            source_object_key=str(blob_ref.get("key") or ""),
            source_etag=str(blob_ref.get("etag") or ""),
            source_size_bytes=int(blob_ref.get("size") or 0),
            source_content_type=str(blob_ref.get("content_type") or source_metadata.get("mime_type") or ""),
            active=True,
            version_ordinal=version_ordinal,
            superseded_at="",
            parser_provenance=bundle.parser_provenance,
            extraction_status=bundle.extraction_status,
            extraction_error=bundle.extraction_error,
            metadata_confidence=float(corpus_metadata.get("metadata_confidence") or 0.0),
            lifecycle_phase=str(corpus_metadata.get("lifecycle_phase") or ""),
            doc_type=str(corpus_metadata.get("doc_type") or ""),
            program_entities=[str(item) for item in list(corpus_metadata.get("program_entities") or []) if str(item)],
            signal_summary=signal_summary,
        )
        stores.doc_store.upsert_document(document_record)
        try:
            stores.chunk_store.add_chunks(chunk_records, tenant_id=tenant_id)
            _sync_requirement_inventory_for_document(
                stores,
                document_record,
                tenant_id=tenant_id,
                chunk_records=chunk_records,
            )
        except Exception:
            stores.doc_store.delete_document(doc_id, tenant_id=tenant_id)
            raise
        _supersede_prior_source_versions(
            stores,
            same_source_records,
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            source_type=source_type,
            source_identity=candidate_identity,
            active_doc_id=doc_id,
            superseded_at=ingested_at,
        )
        if getattr(stores, "graph_store", None) is not None:
            try:
                stores.graph_store.ingest_document(  # type: ignore[union-attr]
                    document_record,
                    chunk_records,
                    tenant_id=tenant_id,
                )
            except Exception as exc:
                logger.warning("Graph ingest failed for %s: %s", doc_id, exc)
        ingested_doc_ids.append(doc_id)
        existing_records.append(document_record)
    return ingested_doc_ids


def ensure_kb_indexed(
    settings: Settings,
    stores: KnowledgeStores,
    tenant_id: str,
    *,
    collection_id: str | None = None,
    attempt_sync: bool | None = None,
) -> KBCoverageStatus:
    effective_collection_id = collection_id or getattr(settings, "default_collection_id", "default")
    status = get_kb_coverage_status(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
    )
    if status.ready or not status.missing_source_paths:
        return status

    should_sync = (
        bool(getattr(settings, "seed_demo_kb_on_startup", True))
        if attempt_sync is None
        else bool(attempt_sync)
    )
    if not should_sync:
        return status

    missing_paths = [Path(path) for path in status.missing_source_paths]
    try:
        synced_doc_ids = ingest_paths(
            settings,
            stores,
            missing_paths,
            source_type="kb",
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
        )
    except Exception as exc:
        return build_kb_coverage_status(
            settings,
            [],
            tenant_id=tenant_id,
            collection_id=effective_collection_id,
            sync_attempted=True,
            sync_error=str(exc),
        )

    refreshed = get_kb_coverage_status(
        settings,
        stores,
        tenant_id=tenant_id,
        collection_id=effective_collection_id,
    )
    return KBCoverageStatus(
        tenant_id=refreshed.tenant_id,
        collection_id=refreshed.collection_id,
        configured_source_paths=refreshed.configured_source_paths,
        missing_source_paths=refreshed.missing_source_paths,
        indexed_source_paths=refreshed.indexed_source_paths,
        indexed_doc_count=refreshed.indexed_doc_count,
        sync_attempted=True,
        sync_error=refreshed.sync_error,
        synced_doc_ids=tuple(str(doc_id) for doc_id in synced_doc_ids if str(doc_id)),
    )
