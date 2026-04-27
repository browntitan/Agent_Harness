from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List

from agentic_chatbot_next.documents.evidence.models import EvidenceBinderResult, EvidenceRow
from agentic_chatbot_next.runtime.artifacts import register_workspace_artifact


_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9_.-]+")


def safe_stem(value: str, *, fallback: str = "evidence_binder") -> str:
    stem = Path(str(value or "")).stem or fallback
    clean = _FILENAME_SAFE_RE.sub("_", stem).strip("._")
    return (clean or fallback)[:72]


def ensure_workspace(session: object) -> object:
    workspace = getattr(session, "workspace", None)
    if workspace is not None:
        return workspace
    workspace_root = str(getattr(session, "workspace_root", "") or "").strip()
    session_id = str(getattr(session, "session_id", "") or "").strip()
    if not workspace_root or not session_id:
        raise ValueError("No session workspace is available.")
    from agentic_chatbot_next.sandbox.workspace import SessionWorkspace

    workspace = SessionWorkspace(session_id=session_id, root=Path(workspace_root))
    workspace.open()
    session.workspace = workspace
    return workspace


def write_binary_artifact(session: object, filename: str, content: bytes) -> Dict[str, Any]:
    workspace = ensure_workspace(session)
    path = workspace.root / filename
    path.write_bytes(content)
    return register_workspace_artifact(session, filename=filename, label=filename)


def evidence_table_csv(rows: Iterable[EvidenceRow]) -> str:
    fields = [
        "Evidence ID",
        "Claim or Generated Statement",
        "Source Excerpt",
        "Source Document",
        "Source Location",
        "Source Hash",
        "Producing Tool/Artifact",
        "Citation Status",
        "Review Status",
        "Notes",
    ]
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fields)
    writer.writeheader()
    for row in rows:
        writer.writerow(
            {
                "Evidence ID": row.evidence_id,
                "Claim or Generated Statement": row.claim,
                "Source Excerpt": row.source_excerpt,
                "Source Document": row.source_document,
                "Source Location": row.source_location,
                "Source Hash": row.source_hash,
                "Producing Tool/Artifact": row.producing_tool or row.artifact_filename,
                "Citation Status": row.citation_status,
                "Review Status": row.review_status,
                "Notes": row.notes,
            }
        )
    return buffer.getvalue()


def source_excerpts_jsonl(result: EvidenceBinderResult) -> str:
    lines = []
    for row in result.evidence_rows:
        if not row.source_excerpt:
            continue
        lines.append(json.dumps(row.to_dict(), sort_keys=True))
    return "\n".join(lines) + ("\n" if lines else "")


def open_issues_markdown(result: EvidenceBinderResult) -> str:
    lines = [
        "# Evidence Binder Open Issues",
        "",
        f"- Binder ID: {result.binder_id}",
        f"- Title: {result.binder_title}",
    ]
    if result.warnings:
        lines.extend(["", "## Warnings"])
        lines.extend(f"- {warning}" for warning in result.warnings)
    if result.open_issues:
        lines.extend(["", "## Issues"])
        for issue in result.open_issues:
            suffix = f" ({issue.related_evidence_id})" if issue.related_evidence_id else ""
            lines.append(f"- [{issue.severity}] {issue.message}{suffix}")
    else:
        lines.extend(["", "No open issues were recorded."])
    return "\n".join(lines).rstrip() + "\n"


def manifest_json(result: EvidenceBinderResult) -> str:
    return json.dumps(result.to_dict(), indent=2, sort_keys=True)


def render_docx_bytes(result: EvidenceBinderResult) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(result.binder_title or "Evidence Binder", level=0)
    if result.objective:
        document.add_paragraph(f"Objective: {result.objective}")
    document.add_paragraph(f"Binder ID: {result.binder_id}")
    document.add_paragraph(f"Generated: {result.generated_at}")
    document.add_paragraph(
        "This binder packages evidence and provenance for review. Classification marking and redaction review are not automatic in V1."
    )

    document.add_heading("Source Inventory", level=1)
    source_table = document.add_table(rows=1, cols=6)
    source_table.style = "Table Grid"
    for index, header in enumerate(["Doc ID", "Title", "Type", "Collection", "Hash", "Parser"]):
        source_table.rows[0].cells[index].text = header
    for source in result.source_documents:
        cells = source_table.add_row().cells
        cells[0].text = source.doc_id
        cells[1].text = source.title
        cells[2].text = source.file_type
        cells[3].text = source.collection_id
        cells[4].text = source.content_hash[:16]
        cells[5].text = source.parser_path

    document.add_heading("Included Artifact Inventory", level=1)
    artifact_table = document.add_table(rows=1, cols=5)
    artifact_table.style = "Table Grid"
    for index, header in enumerate(["Filename", "Artifact Ref", "Type", "Size", "Included"]):
        artifact_table.rows[0].cells[index].text = header
    for artifact in result.artifacts:
        cells = artifact_table.add_row().cells
        cells[0].text = artifact.filename
        cells[1].text = artifact.artifact_ref
        cells[2].text = artifact.content_type
        cells[3].text = str(artifact.size_bytes)
        cells[4].text = "yes" if artifact.included_in_zip else "no"

    document.add_heading("Claim To Evidence Table", level=1)
    evidence_table = document.add_table(rows=1, cols=6)
    evidence_table.style = "Table Grid"
    for index, header in enumerate(["ID", "Claim", "Source", "Location", "Citation", "Review"]):
        evidence_table.rows[0].cells[index].text = header
    for row in result.evidence_rows[:80]:
        cells = evidence_table.add_row().cells
        cells[0].text = row.evidence_id
        cells[1].text = row.claim[:700]
        cells[2].text = row.source_document
        cells[3].text = row.source_location
        cells[4].text = row.citation_status
        cells[5].text = row.review_status

    document.add_heading("Open Issues And Citation Gaps", level=1)
    if result.open_issues:
        for issue in result.open_issues:
            document.add_paragraph(f"[{issue.severity}] {issue.message}", style="List Bullet")
    else:
        document.add_paragraph("No open issues were recorded.")

    document.add_heading("Source Excerpt Appendix", level=1)
    for row in result.evidence_rows[:120]:
        if not row.source_excerpt:
            continue
        document.add_heading(row.evidence_id, level=2)
        document.add_paragraph(row.source_excerpt[:1200])
        if row.source_location:
            document.add_paragraph(f"Location: {row.source_location}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_zip_bytes(
    result: EvidenceBinderResult,
    *,
    workspace_root: Path,
    included_artifact_filenames: List[str],
    original_source_paths: List[Path] | None = None,
) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("evidence_manifest.json", manifest_json(result))
        archive.writestr("evidence_table.csv", evidence_table_csv(result.evidence_rows))
        archive.writestr("source_excerpts.jsonl", source_excerpts_jsonl(result))
        archive.writestr("open_issues.md", open_issues_markdown(result))
        for filename in included_artifact_filenames:
            path = workspace_root / filename
            if path.exists() and path.is_file():
                archive.write(path, arcname=f"included_artifacts/{filename}")
        for source_path in original_source_paths or []:
            if source_path.exists() and source_path.is_file():
                archive.write(source_path, arcname=f"original_sources/{source_path.name}")
    return buffer.getvalue()


__all__ = [
    "build_zip_bytes",
    "evidence_table_csv",
    "manifest_json",
    "open_issues_markdown",
    "render_docx_bytes",
    "safe_stem",
    "source_excerpts_jsonl",
    "write_binary_artifact",
]
