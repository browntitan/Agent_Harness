from __future__ import annotations

import builtins
from pathlib import Path
from types import SimpleNamespace

import pytest
from langchain_core.documents import Document
from openpyxl import Workbook

from agentic_chatbot_next.benchmark.defense_corpus import (
    DefenseBenchmarkQuestion,
    evaluate_defense_contract,
)
from agentic_chatbot_next.contracts.rag import Citation, RagContract, RetrievalSummary
from agentic_chatbot_next.persistence.postgres.chunks import ScoredChunk
from agentic_chatbot_next.rag.adaptive import CorpusRetrievalAdapter
from agentic_chatbot_next.rag.citations import render_citation_location
from agentic_chatbot_next.rag.ingest import _build_chunk_records, _load_documents, ingest_paths
from agentic_chatbot_next.rag.workbook_loader import load_workbook_documents
from agentic_chatbot_next.runtime.task_plan import build_fallback_plan


class _FakeDocStore:
    def __init__(self) -> None:
        self.records = []

    def document_exists(
        self,
        doc_id: str,
        content_hash: str,
        tenant_id: str,
        *,
        collection_id: str = "",
        source_type: str = "",
        title: str = "",
    ) -> bool:
        del doc_id, content_hash, tenant_id, collection_id, source_type, title
        return False

    def upsert_document(self, doc) -> None:
        self.records.append(doc)

    def delete_document(self, doc_id: str, tenant_id: str) -> None:
        del doc_id, tenant_id


class _FakeChunkStore:
    def __init__(self) -> None:
        self.added = []

    def add_chunks(self, chunks, tenant_id: str) -> None:
        self.added.append((tenant_id, list(chunks)))


def _make_ingest_settings(tmp_path: Path):
    return SimpleNamespace(
        object_store_backend="local",
        default_collection_id="default",
        docling_enabled=False,
        ocr_enabled=False,
        ocr_min_page_chars=50,
        ocr_language="en",
        ocr_use_gpu=False,
        chunk_size=240,
        chunk_overlap=40,
    )


def _defense_doc_path(*parts: str) -> Path:
    return Path(__file__).resolve().parents[1].joinpath("defense_rag_test_corpus", "documents", *parts)


@pytest.mark.parametrize(
    ("relative_path", "needle"),
    [
        (("txt", "asterion_issue_digest_draft.txt"), "Asterion"),
        (("pdf", "asterion_monthly_status_review_final.pdf"), "Asterion"),
        (("xlsx", "asterion_budget_schedule_tracker.xlsx"), "Milestone"),
        (("docx", "asterion_ecp_04_rev_c.docx"), "Asterion"),
    ],
)
def test_load_documents_extracts_supported_defense_formats(tmp_path: Path, relative_path: tuple[str, str], needle: str):
    docs = _load_documents(_defense_doc_path(*relative_path), _make_ingest_settings(tmp_path))

    content = "\n".join(doc.page_content for doc in docs)
    assert len(content.strip()) > 50
    assert needle.lower() in content.lower()


def test_docx_loader_reports_clear_parser_failure(tmp_path: Path):
    bad_docx = tmp_path / "broken.docx"
    bad_docx.write_bytes(b"not a valid docx archive")

    with pytest.raises(RuntimeError, match="DOCX extraction failed"):
        _load_documents(bad_docx, _make_ingest_settings(tmp_path))


def test_docx_loader_falls_back_when_docling_enabled_but_missing(tmp_path: Path, monkeypatch):
    from docx import Document as DocxDocument

    docx_path = tmp_path / "fallback.docx"
    doc = DocxDocument()
    doc.add_paragraph("Asterion monthly status reports are due every Friday.")
    doc.save(docx_path)

    settings = _make_ingest_settings(tmp_path)
    settings.docling_enabled = True
    original_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name.startswith("docling"):
            raise ImportError("docling intentionally absent from slim image")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    docs = _load_documents(docx_path, settings)

    assert docs
    assert "monthly status reports" in docs[0].page_content


def test_workbook_loader_emits_sheet_aware_documents_for_defense_corpus():
    workbook_path = (
        Path(__file__).resolve().parents[1]
        / "defense_rag_test_corpus"
        / "documents"
        / "xlsx"
        / "asterion_budget_schedule_tracker.xlsx"
    )

    docs = load_workbook_documents(workbook_path)

    assert docs
    assert any((doc.metadata or {}).get("sheet_name") == "IMS" for doc in docs)
    assert any((doc.metadata or {}).get("chunk_type") == "worksheet_summary" for doc in docs)
    ims_rows = [doc for doc in docs if (doc.metadata or {}).get("sheet_name") == "IMS"]
    assert any("Milestone:" in doc.page_content for doc in ims_rows)
    assert any((doc.metadata or {}).get("row_start") for doc in ims_rows)


def test_ingest_paths_recurses_directories_and_preserves_workbook_provenance(tmp_path: Path):
    docs_root = tmp_path / "documents" / "nested"
    docs_root.mkdir(parents=True)
    (docs_root / "notes.txt").write_text("Asterion supplier note for North Coast Systems LLC.", encoding="utf-8")

    workbook_path = docs_root / "tracker.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "IMS"
    sheet.append(["Milestone", "Current Date", "Variance"])
    sheet.append(["CDR", "2028-09-26", 43])
    workbook.save(workbook_path)

    stores = SimpleNamespace(
        doc_store=_FakeDocStore(),
        chunk_store=_FakeChunkStore(),
        graph_store=None,
    )

    doc_ids = ingest_paths(
        _make_ingest_settings(tmp_path),
        stores,
        [tmp_path / "documents"],
        source_type="kb",
        tenant_id="tenant-a",
        collection_id="defense-rag-test",
    )

    assert len(doc_ids) == 2
    assert all(record.collection_id == "defense-rag-test" for record in stores.doc_store.records)
    all_chunks = [chunk for _, batch in stores.chunk_store.added for chunk in batch]
    assert any(chunk.sheet_name == "IMS" for chunk in all_chunks)
    assert any(chunk.row_start == 2 and chunk.cell_range for chunk in all_chunks)
    assert any((chunk.metadata_json.get("location") or {}).get("sheet_name") == "IMS" for chunk in all_chunks)
    assert all(record.source_metadata.get("index_metadata") for record in stores.doc_store.records)


def test_ingest_paths_allows_same_file_in_multiple_collections(tmp_path: Path):
    doc_path = tmp_path / "regional_spend.csv"
    doc_path.write_text("region,annual_spend_usd\nNA,850000\n", encoding="utf-8")

    class CollectionAwareDocStore:
        def __init__(self) -> None:
            self.records = []

        def document_exists(
            self,
            doc_id: str,
            content_hash: str,
            tenant_id: str,
            *,
            collection_id: str = "",
            source_type: str = "",
            title: str = "",
        ) -> bool:
            return any(
                record.doc_id == doc_id
                or (
                    record.tenant_id == tenant_id
                    and record.content_hash == content_hash
                    and record.collection_id == collection_id
                    and record.source_type == source_type
                    and record.title == title
                )
                for record in self.records
            )

        def upsert_document(self, doc) -> None:
            self.records.append(doc)

        def delete_document(self, doc_id: str, tenant_id: str) -> None:
            self.records = [
                record for record in self.records if not (record.doc_id == doc_id and record.tenant_id == tenant_id)
            ]

    stores = SimpleNamespace(
        doc_store=CollectionAwareDocStore(),
        chunk_store=_FakeChunkStore(),
        graph_store=None,
    )

    first_ids = ingest_paths(
        _make_ingest_settings(tmp_path),
        stores,
        [doc_path],
        source_type="upload",
        tenant_id="tenant-a",
        collection_id="collection-a",
    )
    second_ids = ingest_paths(
        _make_ingest_settings(tmp_path),
        stores,
        [doc_path],
        source_type="upload",
        tenant_id="tenant-a",
        collection_id="collection-b",
    )

    assert len(first_ids) == 1
    assert len(second_ids) == 1
    assert first_ids[0] != second_ids[0]
    assert {record.collection_id for record in stores.doc_store.records} == {"collection-a", "collection-b"}


def test_build_chunk_records_normalizes_duplicate_chunk_indexes():
    chunks = [
        Document(page_content="alpha", metadata={"chunk_index": 0}),
        Document(page_content="beta", metadata={"chunk_index": 0}),
        Document(page_content="gamma", metadata={"chunk_index": "1"}),
    ]

    records = _build_chunk_records(chunks, "doc-1", collection_id="defense-rag-test")

    assert [record.chunk_index for record in records] == [0, 1, 2]
    assert len({record.chunk_id for record in records}) == 3


def test_render_citation_location_prefers_workbook_sheet_and_rows():
    location = render_citation_location(
        {
            "sheet_name": "Scorecard",
            "row_start": 5,
            "row_end": 5,
            "cell_range": "A5:E5",
        }
    )

    assert location == "Scorecard row 5 A5:E5"


def test_query_heuristics_prefer_final_workbook_for_latest_schedule_question():
    adapter = CorpusRetrievalAdapter(
        SimpleNamespace(doc_store=SimpleNamespace(), graph_store=None, chunk_store=SimpleNamespace()),
        settings=SimpleNamespace(graph_search_enabled=False, default_collection_id="default"),
        session=SimpleNamespace(tenant_id="tenant-a", metadata={"collection_id": "defense-rag-test"}),
    )

    draft_chunk = ScoredChunk(
        doc=Document(
            page_content="Draft review discussing tentative schedule movement.",
            metadata={
                "chunk_id": "draft#chunk0001",
                "doc_id": "draft",
                "title": "asterion_monthly_status_review_draft.pdf",
                "file_type": "pdf",
            },
        ),
        score=0.82,
        method="keyword",
    )
    final_chunk = ScoredChunk(
        doc=Document(
            page_content="Workbook entry with current approved CDR date and variance.",
            metadata={
                "chunk_id": "final#chunk0001",
                "doc_id": "final",
                "title": "asterion_budget_schedule_tracker.xlsx",
                "file_type": "xlsx",
                "sheet_name": "IMS",
            },
        ),
        score=0.75,
        method="vector",
    )

    reranked = adapter._rerank_with_query_heuristics(
        "What is the approved current CDR date for Asterion?",
        [draft_chunk, final_chunk],
    )

    assert reranked[0].doc.metadata["doc_id"] == "final"


def test_defense_benchmark_evaluation_flags_missing_workbook_citation():
    question = DefenseBenchmarkQuestion(
        question_id="E02",
        difficulty="Easy",
        question_text="What is the approved current CDR date for Asterion?",
        expected_answer="26 Sep 2028.",
        source_documents=(
            "asterion_budget_schedule_tracker.xlsx",
            "asterion_monthly_status_review_final.pdf",
        ),
        supporting_references=("IMS!C3", "Section 1 Executive Summary"),
    )
    contract = RagContract(
        answer="The approved current CDR date is 26 Sep 2028.",
        citations=[
            Citation(
                citation_id="c1",
                doc_id="doc-final",
                title="asterion_monthly_status_review_final.pdf",
                source_type="kb",
                location="page 1",
                snippet="Current approved CDR date is 26 Sep 2028.",
            )
        ],
        used_citation_ids=["c1"],
        confidence=0.8,
        retrieval_summary=RetrievalSummary(query_used=question.question_text),
    )

    result = evaluate_defense_contract(question, contract)

    assert result.answer_correct is True
    assert result.citation_source_match is True
    assert "workbook_not_retrieved" in result.diagnostics


def test_document_research_campaign_adds_workbook_followup_task():
    tasks = build_fallback_plan(
        "Identify all documents with current schedule variance and budget impacts across the corpus.",
        max_tasks=4,
    )

    assert [task["executor"] for task in tasks[:3]] == ["rag_worker", "data_analyst", "rag_worker"]
    assert tasks[0]["produces_artifacts"] == ["doc_focus"]
    assert "doc_focus" in tasks[1]["consumes_artifacts"]
    assert "analysis_summary" in tasks[2]["consumes_artifacts"]
